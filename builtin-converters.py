#!/usr/bin/env python3
"""
Built-in Format Converters

This module contains built-in converters for common file formats.
"""

import os
import logging
import json
import csv
import shutil
import subprocess
from typing import List, Dict, Any, Optional, Tuple

logger = logging.getLogger(__name__)

# Import the register_converter function from the main module
try:
    from format_converter import register_converter
except ImportError:
    # If running as standalone, define a dummy function
    def register_converter(source_format, target_format, converter_func):
        """Dummy register_converter function."""
        pass

# -------------------------------------------------------------------------
# Utility functions
# -------------------------------------------------------------------------

def ensure_package_installed(package_name: str) -> bool:
    """
    Ensure a Python package is installed.
    
    Args:
        package_name: Name of the package to check/install
        
    Returns:
        True if package is available (installed or just now installed), False otherwise
    """
    try:
        __import__(package_name)
        return True
    except ImportError:
        logger.warning(f"Package {package_name} is not installed. Some conversions may fail.")
        return False

# -------------------------------------------------------------------------
# Document format converters
# -------------------------------------------------------------------------

def docx_to_pdf(input_path: str, output_path: str) -> None:
    """
    Convert DOCX to PDF.
    
    Args:
        input_path: Path to the input DOCX file
        output_path: Path to save the converted PDF file
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('docx') or not ensure_package_installed('PyPDF2'):
        raise ImportError("Required packages not installed: python-docx, PyPDF2")
    
    try:
        from docx import Document
        from PyPDF2 import PdfWriter
        import io
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Open the Word document
        doc = Document(input_path)
        
        # Create a PDF writer
        pdf_writer = PdfWriter()
        
        # Process each paragraph in the document
        for para in doc.paragraphs:
            # Create a PDF in memory
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=letter)
            c.drawString(72, 800, para.text)  # Position text at (72, 800)
            c.save()
            
            # Move to the beginning of the StringIO buffer
            packet.seek(0)
            
            # Create a PDF reader object
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(packet)
            
            # Add the PDF page
            pdf_writer.add_page(pdf_reader.pages[0])
        
        # Write the PDF to disk
        with open(output_path, 'wb') as f:
            pdf_writer.write(f)
            
    except Exception as e:
        logger.error(f"Error converting DOCX to PDF: {str(e)}")
        raise

def docx_to_txt(input_path: str, output_path: str) -> None:
    """
    Convert DOCX to plain text.
    
    Args:
        input_path: Path to the input DOCX file
        output_path: Path to save the converted TXT file
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('docx'):
        raise ImportError("Required package not installed: python-docx")
    
    try:
        from docx import Document
        
        # Open the Word document
        doc = Document(input_path)
        
        # Extract text from paragraphs
        text = "\n\n".join([para.text for para in doc.paragraphs])
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
            
    except Exception as e:
        logger.error(f"Error converting DOCX to TXT: {str(e)}")
        raise

def txt_to_pdf(input_path: str, output_path: str) -> None:
    """
    Convert TXT to PDF.
    
    Args:
        input_path: Path to the input TXT file
        output_path: Path to save the converted PDF file
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        
        # Read the text file
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Create a PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        
        # Create a style
        styles = getSampleStyleSheet()
        style = styles["Normal"]
        
        # Create paragraphs from the text
        paragraphs = []
        for line in text.split('\n'):
            if line.strip():  # Skip empty lines
                paragraphs.append(Paragraph(line, style))
            else:
                # Add an empty paragraph for blank lines
                paragraphs.append(Paragraph("<br/>", style))
        
        # Build the PDF
        doc.build(paragraphs)
            
    except ImportError:
        logger.error("Required package not installed: reportlab")
        raise
    except Exception as e:
        logger.error(f"Error converting TXT to PDF: {str(e)}")
        raise

def pdf_to_txt(input_path: str, output_path: str) -> None:
    """
    Extract text from PDF.
    
    Args:
        input_path: Path to the input PDF file
        output_path: Path to save the extracted text
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If extraction fails
    """
    if not ensure_package_installed('PyPDF2'):
        raise ImportError("Required package not installed: PyPDF2")
    
    try:
        from PyPDF2 import PdfReader
        
        # Open the PDF
        reader = PdfReader(input_path)
        
        # Extract text from each page
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
            
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def html_to_txt(input_path: str, output_path: str) -> None:
    """
    Convert HTML to plain text.
    
    Args:
        input_path: Path to the input HTML file
        output_path: Path to save the extracted text
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If extraction fails
    """
    if not ensure_package_installed('lxml'):
        raise ImportError("Required package not installed: lxml")
    
    try:
        from lxml import html
        
        # Read the HTML file
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse the HTML
        tree = html.fromstring(html_content)
        
        # Extract text (remove scripts, style elements, etc.)
        for element in tree.xpath('//script | //style'):
            element.getparent().remove(element)
        
        # Get the text content
        text = tree.text_content()
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
            
    except Exception as e:
        logger.error(f"Error converting HTML to TXT: {str(e)}")
        raise

def markdown_to_html(input_path: str, output_path: str) -> None:
    """
    Convert Markdown to HTML.
    
    Args:
        input_path: Path to the input Markdown file
        output_path: Path to save the converted HTML
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('markdown'):
        raise ImportError("Required package not installed: markdown")
    
    try:
        import markdown
        
        # Read the Markdown file
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert to HTML
        html_content = markdown.markdown(md_content)
        
        # Wrap with basic HTML structure
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted from {os.path.basename(input_path)}</title>
</head>
<body>
{html_content}
</body>
</html>
"""
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
            
    except Exception as e:
        logger.error(f"Error converting Markdown to HTML: {str(e)}")
        raise

def html_to_markdown(input_path: str, output_path: str) -> None:
    """
    Convert HTML to Markdown.
    
    Args:
        input_path: Path to the input HTML file
        output_path: Path to save the converted Markdown
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    try:
        # Try to import html2text
        import html2text
        
        # Read the HTML file
        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Convert to Markdown
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.escape_snob = True
        md_content = h.handle(html_content)
        
        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
            
    except ImportError:
        logger.error("Required package not installed: html2text")
        raise
    except Exception as e:
        logger.error(f"Error converting HTML to Markdown: {str(e)}")
        raise

# -------------------------------------------------------------------------
# Spreadsheet format converters
# -------------------------------------------------------------------------

def xlsx_to_csv(input_path: str, output_path: str) -> None:
    """
    Convert XLSX to CSV.
    
    Args:
        input_path: Path to the input XLSX file
        output_path: Path to save the converted CSV
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('openpyxl'):
        raise ImportError("Required package not installed: openpyxl")
    
    try:
        from openpyxl import load_workbook
        
        # Open the workbook
        wb = load_workbook(filename=input_path)
        
        # Get the active worksheet
        ws = wb.active
        
        # Write to CSV
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(row)
                
    except Exception as e:
        logger.error(f"Error converting XLSX to CSV: {str(e)}")
        raise

def csv_to_xlsx(input_path: str, output_path: str) -> None:
    """
    Convert CSV to XLSX.
    
    Args:
        input_path: Path to the input CSV file
        output_path: Path to save the converted XLSX
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('openpyxl'):
        raise ImportError("Required package not installed: openpyxl")
    
    try:
        from openpyxl import Workbook
        
        # Create a new workbook
        wb = Workbook()
        ws = wb.active
        
        # Read the CSV file
        with open(input_path, 'r', newline='', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                ws.append(row)
        
        # Save the workbook
        wb.save(output_path)
                
    except Exception as e:
        logger.error(f"Error converting CSV to XLSX: {str(e)}")
        raise

def csv_to_json(input_path: str, output_path: str) -> None:
    """
    Convert CSV to JSON.
    
    Args:
        input_path: Path to the input CSV file
        output_path: Path to save the converted JSON
    
    Raises:
        Exception: If conversion fails
    """
    try:
        # Read the CSV file
        with open(input_path, 'r', newline='', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            data = list(reader)
        
        # Write to JSON
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4)
                
    except Exception as e:
        logger.error(f"Error converting CSV to JSON: {str(e)}")
        raise

def json_to_csv(input_path: str, output_path: str) -> None:
    """
    Convert JSON to CSV.
    
    Args:
        input_path: Path to the input JSON file
        output_path: Path to save the converted CSV
    
    Raises:
        Exception: If conversion fails
    """
    try:
        # Read the JSON file
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Make sure data is a list of dictionaries
        if not isinstance(data, list):
            data = [data]
        
        # Get all unique keys
        all_keys = set()
        for item in data:
            if isinstance(item, dict):
                all_keys.update(item.keys())
        
        # Sort keys for consistent output
        fieldnames = sorted(all_keys)
        
        # Write to CSV
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            for item in data:
                if isinstance(item, dict):
                    writer.writerow(item)
                
    except Exception as e:
        logger.error(f"Error converting JSON to CSV: {str(e)}")
        raise

# -------------------------------------------------------------------------
# Text and markup format converters
# -------------------------------------------------------------------------

def json_to_xml(input_path: str, output_path: str) -> None:
    """
    Convert JSON to XML.
    
    Args:
        input_path: Path to the input JSON file
        output_path: Path to save the converted XML
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('lxml'):
        raise ImportError("Required package not installed: lxml")
    
    try:
        from lxml import etree
        
        # Read the JSON file
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Helper function to convert JSON to XML
        def json_to_xml_element(parent, key, value):
            """Convert JSON to XML recursively."""
            if isinstance(value, dict):
                # For dictionaries, create a new element with the key as tag
                element = etree.SubElement(parent, key)
                for k, v in value.items():
                    json_to_xml_element(element, k, v)
            elif isinstance(value, list):
                # For lists, create an item element for each item
                for item in value:
                    if isinstance(item, dict):
                        # If item is a dict, use key as element name
                        element = etree.SubElement(parent, key)
                        for k, v in item.items():
                            json_to_xml_element(element, k, v)
                    else:
                        # If item is a primitive, use key as element name
                        element = etree.SubElement(parent, key)
                        element.text = str(item)
            else:
                # For primitives, create an element with text content
                element = etree.SubElement(parent, key)
                element.text = str(value)
        
        # Create the root element
        root = etree.Element("root")
        
        # Convert JSON to XML
        if isinstance(data, dict):
            for key, value in data.items():
                json_to_xml_element(root, key, value)
        elif isinstance(data, list):
            for i, item in enumerate(data):
                json_to_xml_element(root, f"item{i}", item)
        else:
            root.text = str(data)
        
        # Create the XML tree
        tree = etree.ElementTree(root)
        
        # Write to XML file with pretty printing
        tree.write(output_path, pretty_print=True, xml_declaration=True, encoding='utf-8')
                
    except Exception as e:
        logger.error(f"Error converting JSON to XML: {str(e)}")
        raise

def xml_to_json(input_path: str, output_path: str) -> None:
    """
    Convert XML to JSON.
    
    Args:
        input_path: Path to the input XML file
        output_path: Path to save the converted JSON
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('lxml'):
        raise ImportError("Required package not installed: lxml")
    
    try:
        from lxml import etree
        
        # Parse the XML file
        tree = etree.parse(input_path)
        root = tree.getroot()
        
        # Helper function to convert XML to JSON
        def xml_to_json_dict(element):
            """Convert XML to JSON recursively."""
            result = {}
            
            # Add attributes
            for name, value in element.attrib.items():
                result[f"@{name}"] = value
            
            # Add child elements
            for child in element:
                child_dict = xml_to_json_dict(child)
                
                if child.tag in result:
                    # If tag already exists, convert to list or append
                    if not isinstance(result[child.tag], list):
                        result[child.tag] = [result[child.tag]]
                    result[child.tag].append(child_dict)
                else:
                    result[child.tag] = child_dict
            
            # Add text content if no children and not empty
            if len(result) == 0 and element.text and element.text.strip():
                return element.text.strip()
            elif len(result) == 0:
                return None
            elif element.text and element.text.strip():
                result["#text"] = element.text.strip()
            
            return result
        
        # Convert XML to JSON
        json_data = {root.tag: xml_to_json_dict(root)}
        
        # Write to JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2)
                
    except Exception as e:
        logger.error(f"Error converting XML to JSON: {str(e)}")
        raise

def json_to_yaml(input_path: str, output_path: str) -> None:
    """
    Convert JSON to YAML.
    
    Args:
        input_path: Path to the input JSON file
        output_path: Path to save the converted YAML
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('yaml'):
        raise ImportError("Required package not installed: PyYAML")
    
    try:
        import yaml
        
        # Read the JSON file
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Write to YAML file
        with open(output_path, 'w', encoding='utf-8') as f:
            yaml.dump(data, f, default_flow_style=False, sort_keys=False)
                
    except Exception as e:
        logger.error(f"Error converting JSON to YAML: {str(e)}")
        raise

def yaml_to_json(input_path: str, output_path: str) -> None:
    """
    Convert YAML to JSON.
    
    Args:
        input_path: Path to the input YAML file
        output_path: Path to save the converted JSON
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('yaml'):
        raise ImportError("Required package not installed: PyYAML")
    
    try:
        import yaml
        
        # Read the YAML file
        with open(input_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        # Write to JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
                
    except Exception as e:
        logger.error(f"Error converting YAML to JSON: {str(e)}")
        raise

# -------------------------------------------------------------------------
# Image format converters
# -------------------------------------------------------------------------

def convert_image(input_path: str, output_path: str) -> None:
    """
    Convert an image from one format to another using Pillow.
    
    Args:
        input_path: Path to the input image
        output_path: Path to save the converted image
    
    Raises:
        ImportError: If required libraries are not installed
        Exception: If conversion fails
    """
    if not ensure_package_installed('PIL'):
        raise ImportError("Required package not installed: Pillow")
    
    try:
        from PIL import Image
        
        # Open the input image
        with Image.open(input_path) as img:
            # If image is in CMYK mode (common in some JPEG files), convert to RGB
            if img.mode == 'CMYK':
                img = img.convert('RGB')
            
            # Save with the target format determined from the file extension
            img.save(output_path)
                
    except Exception as e:
        logger.error(f"Error converting image: {str(e)}")
        raise

# -------------------------------------------------------------------------
# Archive format converters
# -------------------------------------------------------------------------

def create_archive(input_path: str, output_path: str, archive_format: str) -> None:
    """
    Create an archive from a directory.
    
    Args:
        input_path: Path to the input directory
        output_path: Path to save the archive
        archive_format: Format of the archive (zip, tar, etc.)
    
    Raises:
        Exception: If archive creation fails
    """
    try:
        # Make sure input_path is a directory
        if not os.path.isdir(input_path):
            raise ValueError(f"Input path is not a directory: {input_path}")
        
        if archive_format == 'zip':
            shutil.make_archive(
                os.path.splitext(output_path)[0],  # Base name (without extension)
                'zip',  # Format
                input_path  # Root directory
            )
        elif archive_format in ('tar', 'gztar', 'bztar'):
            shutil.make_archive(
                os.path.splitext(output_path)[0],  # Base name (without extension)
                archive_format,  # Format
                input_path  # Root directory
            )
        else:
            raise ValueError(f"Unsupported archive format: {archive_format}")
                
    except Exception as e:
        logger.error(f"Error creating archive: {str(e)}")
        raise

def extract_archive(input_path: str, output_path: str) -> None:
    """
    Extract an archive to a directory.
    
    Args:
        input_path: Path to the input archive
        output_path: Path to extract the archive to
    
    Raises:
        Exception: If archive extraction fails
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs(output_path, exist_ok=True)
        
        # Determine archive format from file extension
        ext = os.path.splitext(input_path)[1].lower()
        
        if ext == '.zip':
            shutil.unpack_archive(input_path, output_path, 'zip')
        elif ext == '.tar':
            shutil.unpack_archive(input_path, output_path, 'tar')
        elif ext in ('.gz', '.tgz'):
            shutil.unpack_archive(input_path, output_path, 'gztar')
        elif ext in ('.bz2', '.tbz2'):
            shutil.unpack_archive(input_path, output_path, 'bztar')
        else:
            raise ValueError(f"Unsupported archive format: {ext}")
                
    except Exception as e:
        logger.error(f"Error extracting archive: {str(e)}")
        raise

# -------------------------------------------------------------------------
# Register converters
# -------------------------------------------------------------------------

def register_builtin_converters():
    """Register all built-in converters."""
    # Document format converters
    register_converter('docx', 'pdf', docx_to_pdf)
    register_converter('docx', 'txt', docx_to_txt)
    register_converter('txt', 'pdf', txt_to_pdf)
    register_converter('pdf', 'txt', pdf_to_txt)
    register_converter('html', 'txt', html_to_txt)
    register_converter('md', 'html', markdown_to_html)
    register_converter('html', 'md', html_to_markdown)
    
    # Spreadsheet format converters
    register_converter('xlsx', 'csv', xlsx_to_csv)
    register_converter('csv', 'xlsx', csv_to_xlsx)
    register_converter('csv', 'json', csv_to_json)
    register_converter('json', 'csv', json_to_csv)
    
    # Text and markup format converters
    register_converter('json', 'xml', json_to_xml)
    register_converter('xml', 'json', xml_to_json)
    register_converter('json', 'yaml', json_to_yaml)
    register_converter('yaml', 'json', yaml_to_json)
    register_converter('yml', 'json', yaml_to_json)  # yml alias
    register_converter('json', 'yml', json_to_yaml)  # yml alias
    
    # Image format converters - register all combinations
    image_formats = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp']
    for src_fmt in image_formats:
        for tgt_fmt in image_formats:
            if src_fmt != tgt_fmt:
                register_converter(src_fmt, tgt_fmt, convert_image)

# If running as standalone, register converters
if __name__ == "__main__":
    register_builtin_converters()

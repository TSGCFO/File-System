"""
Spreadsheet format converters for FileConverter.

This module provides converters for spreadsheet formats, including:
- Microsoft Excel (.xls, .xlsx)
- CSV (.csv)
- TSV (.tsv)
- JSON (.json)
- XML (.xml)
- HTML (.html)
"""

import csv
import json
import tempfile
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from fileconverter.core.registry import BaseConverter
from fileconverter.utils.error_handling import ConversionError
from fileconverter.utils.file_utils import get_file_extension, guess_encoding
from fileconverter.utils.logging_utils import get_logger
from fileconverter.utils.validation import validate_file_path

logger = get_logger(__name__)

# Define supported formats
SUPPORTED_FORMATS = ["xlsx", "xls", "csv", "tsv", "json", "xml", "html", "docx", "pdf", "txt", "md", "yaml", "ini", "toml"]


class SpreadsheetConverter(BaseConverter):
    """Converter for spreadsheet formats."""
    
    @classmethod
    def get_input_formats(cls) -> List[str]:
        """Get the list of input formats supported by this converter."""
        return ["xlsx", "xls", "csv", "tsv", "json"]
    
    @classmethod
    def get_output_formats(cls) -> List[str]:
        """Get the list of output formats supported by this converter."""
        return ["xlsx", "csv", "tsv", "json", "html", "xml", "docx", "pdf", "txt", "md", "yaml", "ini", "toml"]
    
    @classmethod
    def get_format_extensions(cls, format_name: str) -> List[str]:
        """Get the list of file extensions for a specific format."""
        format_map = {
            "xlsx": ["xlsx"],
            "xls": ["xls"],
            "csv": ["csv"],
            "tsv": ["tsv"],
            "json": ["json"],
            "xml": ["xml"],
            "html": ["html", "htm"],
            "docx": ["docx"],
            "pdf": ["pdf"],
            "txt": ["txt"],
            "md": ["md", "markdown"],
            "yaml": ["yaml", "yml"],
            "ini": ["ini", "conf", "cfg"],
            "toml": ["toml"],
        }
        return format_map.get(format_name.lower(), [])
    
    def convert(
        self, 
        input_path: Union[str, Path], 
        output_path: Union[str, Path],
        temp_dir: Union[str, Path],
        parameters: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Convert a spreadsheet from one format to another.
        
        Args:
            input_path: Path to the input file.
            output_path: Path where the output file will be saved.
            temp_dir: Directory for temporary files.
            parameters: Conversion parameters.
        
        Returns:
            Dictionary with information about the conversion.
        
        Raises:
            ConversionError: If the conversion fails.
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        temp_dir = Path(temp_dir)
        
        # Validate paths
        validate_file_path(input_path, must_exist=True)
        
        # Get formats from file extensions
        input_ext = get_file_extension(input_path).lower()
        output_ext = get_file_extension(output_path).lower()
        
        # Map extensions to formats
        input_format = input_ext
        output_format = output_ext
        
        # Verify supported formats
        if input_format not in self.get_input_formats():
            raise ConversionError(f"Unsupported input format: {input_format}")
        
        if output_format not in self.get_output_formats():
            raise ConversionError(f"Unsupported output format: {output_format}")
        
        # Determine conversion method based on input and output formats
        try:
            # Most conversions will use pandas
            import pandas as pd
            
            # Load spreadsheet data based on input format
            if input_format in ["xlsx", "xls"]:
                df = self._load_excel(input_path, parameters)
            elif input_format == "csv":
                df = self._load_csv(input_path, parameters)
            elif input_format == "tsv":
                df = self._load_tsv(input_path, parameters)
            elif input_format == "json":
                df = self._load_json(input_path, parameters)
            else:
                raise ConversionError(f"Unsupported input format: {input_format}")
            
            # Save spreadsheet data based on output format
            if output_format in ["xlsx", "xls"]:
                self._save_excel(df, output_path, parameters)
            elif output_format == "csv":
                self._save_csv(df, output_path, parameters)
            elif output_format == "tsv":
                self._save_tsv(df, output_path, parameters)
            elif output_format == "json":
                self._save_json(df, output_path, parameters)
            elif output_format == "html":
                self._save_html(df, output_path, parameters)
            elif output_format == "xml":
                self._save_xml(df, output_path, parameters)
            elif output_format == "docx":
                self._save_docx(df, output_path, parameters)
            elif output_format == "pdf":
                self._save_pdf(df, output_path, parameters)
            elif output_format == "txt":
                self._save_txt(df, output_path, parameters)
            elif output_format == "md":
                self._save_md(df, output_path, parameters)
            elif output_format == "yaml":
                self._save_yaml(df, output_path, parameters)
            elif output_format == "ini":
                self._save_ini(df, output_path, parameters)
            elif output_format == "toml":
                self._save_toml(df, output_path, parameters)
            else:
                raise ConversionError(f"Unsupported output format: {output_format}")
        
        except ImportError as e:
            raise ConversionError(
                f"Missing required dependency: {str(e)}. "
                "Please install pandas with 'pip install pandas'."
            )
        except Exception as e:
            raise ConversionError(
                f"Failed to convert {input_format} to {output_format}: {str(e)}"
            )
        
        return {
            "input_format": input_format,
            "output_format": output_format,
            "input_path": str(input_path),
            "output_path": str(output_path),
            "rows_processed": len(df) if 'df' in locals() else None,
        }
    
    def get_parameters(self) -> Dict[str, Dict[str, Any]]:
        """Get the parameters supported by this converter."""
        return {
            "excel": {
                "sheet_name": {
                    "type": "string",
                    "description": "Sheet name to read/write (excel only)",
                    "default": 0,  # 0 = first sheet
                },
                "header": {
                    "type": "boolean",
                    "description": "Whether the file has a header row",
                    "default": True,
                },
                "index": {
                    "type": "boolean",
                    "description": "Whether to include index in output",
                    "default": False,
                },
            },
            "csv": {
                "delimiter": {
                    "type": "string",
                    "description": "Delimiter character",
                    "default": ",",
                },
                "quotechar": {
                    "type": "string",
                    "description": "Quote character",
                    "default": '"',
                },
                "encoding": {
                    "type": "string",
                    "description": "File encoding",
                    "default": "utf-8",
                },
                "header": {
                    "type": "boolean",
                    "description": "Whether the file has a header row",
                    "default": True,
                },
            },
            "json": {
                "orient": {
                    "type": "string",
                    "description": "JSON orientation",
                    "default": "records",
                    "options": ["records", "columns", "index", "split", "table"],
                },
                "indent": {
                    "type": "integer",
                    "description": "Indentation level",
                    "default": 2,
                },
            },
            "html": {
                "table_id": {
                    "type": "string",
                    "description": "HTML table ID",
                    "default": "data",
                },
                "index": {
                    "type": "boolean",
                    "description": "Whether to include index in output",
                    "default": False,
                },
                "classes": {
                    "type": "string",
                    "description": "CSS classes for the table",
                    "default": "dataframe",
                },
                "escape": {
                    "type": "boolean",
                    "description": "Whether to escape HTML entities",
                    "default": True,
                },
            },
            "xml": {
                "root": {
                    "type": "string",
                    "description": "Root element name",
                    "default": "data",
                },
                "row": {
                    "type": "string",
                    "description": "Row element name",
                    "default": "row",
                },
            },
            "docx": {
                "title": {
                    "type": "string",
                    "description": "Document title",
                    "default": "Spreadsheet Data",
                },
                "table_style": {
                    "type": "string",
                    "description": "Style for the table",
                    "default": "Table Grid",
                },
            },
            "pdf": {
                "page_size": {
                    "type": "string",
                    "description": "Page size (e.g., A4, Letter)",
                    "default": "A4",
                    "options": ["A4", "Letter", "Legal"],
                },
                "orientation": {
                    "type": "string",
                    "description": "Page orientation",
                    "default": "portrait",
                    "options": ["portrait", "landscape"],
                },
                "margin": {
                    "type": "number",
                    "description": "Page margin in inches",
                    "default": 1.0,
                    "min": 0.0,
                    "max": 3.0,
                },
                "title": {
                    "type": "string",
                    "description": "Document title",
                    "default": "Spreadsheet Data",
                },
            },
        }
    
    def _load_excel(self, file_path: Path, parameters: Dict[str, Any]) -> "pd.DataFrame":
        """Load data from an Excel file.
        
        Args:
            file_path: Path to the Excel file.
            parameters: Conversion parameters.
        
        Returns:
            Pandas DataFrame with the loaded data.
        
        Raises:
            ConversionError: If the file cannot be loaded.
        """
        import pandas as pd
        
        sheet_name = parameters.get("sheet_name", 0)
        header = 0 if parameters.get("header", True) else None
        
        try:
            return pd.read_excel(
                file_path,
                sheet_name=sheet_name,
                header=header
            )
        except Exception as e:
            raise ConversionError(f"Failed to load Excel file: {str(e)}")
    
    def _load_csv(self, file_path: Path, parameters: Dict[str, Any]) -> "pd.DataFrame":
        """Load data from a CSV file.
        
        Args:
            file_path: Path to the CSV file.
            parameters: Conversion parameters.
        
        Returns:
            Pandas DataFrame with the loaded data.
        
        Raises:
            ConversionError: If the file cannot be loaded.
        """
        import pandas as pd
        
        delimiter = parameters.get("delimiter", ",")
        quotechar = parameters.get("quotechar", '"')
        encoding = parameters.get("encoding")
        header = 0 if parameters.get("header", True) else None
        
        if not encoding:
            encoding = guess_encoding(file_path)
        
        try:
            return pd.read_csv(
                file_path,
                delimiter=delimiter,
                quotechar=quotechar,
                encoding=encoding,
                header=header
            )
        except Exception as e:
            raise ConversionError(f"Failed to load CSV file: {str(e)}")
    
    def _load_tsv(self, file_path: Path, parameters: Dict[str, Any]) -> "pd.DataFrame":
        """Load data from a TSV file.
        
        Args:
            file_path: Path to the TSV file.
            parameters: Conversion parameters.
        
        Returns:
            Pandas DataFrame with the loaded data.
        
        Raises:
            ConversionError: If the file cannot be loaded.
        """
        # Use CSV loader with tab delimiter
        parameters["delimiter"] = "\t"
        return self._load_csv(file_path, parameters)
    
    def _load_json(self, file_path: Path, parameters: Dict[str, Any]) -> "pd.DataFrame":
        """Load data from a JSON file.
        
        Args:
            file_path: Path to the JSON file.
            parameters: Conversion parameters.
        
        Returns:
            Pandas DataFrame with the loaded data.
        
        Raises:
            ConversionError: If the file cannot be loaded.
        """
        import pandas as pd
        
        orient = parameters.get("orient", "records")
        encoding = parameters.get("encoding", "utf-8")
        
        try:
            return pd.read_json(
                file_path,
                orient=orient,
                encoding=encoding
            )
        except Exception as e:
            raise ConversionError(f"Failed to load JSON file: {str(e)}")
    
    def _save_excel(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to an Excel file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        sheet_name = parameters.get("sheet_name", "Sheet1")
        index = parameters.get("index", False)
        
        try:
            df.to_excel(
                file_path,
                sheet_name=sheet_name,
                index=index
            )
        except Exception as e:
            raise ConversionError(f"Failed to save Excel file: {str(e)}")
    
    def _save_csv(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a CSV file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        delimiter = parameters.get("delimiter", ",")
        quotechar = parameters.get("quotechar", '"')
        encoding = parameters.get("encoding", "utf-8")
        index = parameters.get("index", False)
        header = parameters.get("header", True)
        
        try:
            df.to_csv(
                file_path,
                sep=delimiter,
                quotechar=quotechar,
                encoding=encoding,
                index=index,
                header=header
            )
        except Exception as e:
            raise ConversionError(f"Failed to save CSV file: {str(e)}")
    
    def _save_tsv(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a TSV file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        # Use CSV saver with tab delimiter
        parameters["delimiter"] = "\t"
        return self._save_csv(df, file_path, parameters)
    
    def _save_json(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a JSON file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        orient = parameters.get("orient", "records")
        indent = parameters.get("indent", 2)
        
        try:
            df.to_json(
                file_path,
                orient=orient,
                indent=indent
            )
        except Exception as e:
            raise ConversionError(f"Failed to save JSON file: {str(e)}")
    
    def _save_html(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to an HTML file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        table_id = parameters.get("table_id", "data")
        index = parameters.get("index", False)
        classes = parameters.get("classes", "dataframe")
        escape = parameters.get("escape", True)
        
        try:
            # Generate HTML table
            html_table = df.to_html(
                index=index,
                table_id=table_id,
                classes=classes,
                escape=escape
            )
            
            # Create full HTML document
            html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Data Table</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        table { border-collapse: collapse; width: 100%; margin-top: 20px; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; }
    </style>
</head>
<body>
    <h1>Data Table</h1>
""" + html_table + """
</body>
</html>"""
            
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_content)
        except Exception as e:
            raise ConversionError(f"Failed to save HTML file: {str(e)}")
    
    def _save_xml(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to an XML file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        root_name = parameters.get("root", "data")
        row_name = parameters.get("row", "row")
        
        try:
            from dicttoxml import dicttoxml
            
            # Convert DataFrame to list of dictionaries
            data = df.to_dict(orient="records")
            
            # Convert to XML
            xml = dicttoxml(
                data,
                custom_root=root_name,
                item_func=lambda x: row_name,
                attr_type=False
            )
            
            # Write to file
            with open(file_path, "wb") as f:
                f.write(xml)
        
        except ImportError:
            # Fallback if dicttoxml is not available
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(f'<?xml version="1.0" encoding="UTF-8"?>\n')
                    f.write(f'<{root_name}>\n')
                    
                    for _, row in df.iterrows():
                        f.write(f'  <{row_name}>\n')
                        for col, value in row.items():
                            if pd.isna(value):
                                f.write(f'    <{col}/>\n')
                            else:
                                f.write(f'    <{col}>{value}</{col}>\n')
                        f.write(f'  </{row_name}>\n')
                    
                    f.write(f'</{root_name}>\n')
            
            except Exception as e:
                raise ConversionError(f"Failed to save XML file: {str(e)}")
        
        except Exception as e:
            raise ConversionError(f"Failed to save XML file: {str(e)}")
    
    def _save_docx(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a DOCX file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        try:
            import docx
            from docx import Document
            from docx.shared import Inches
            
            document = Document()
            
            # Add title
            title = parameters.get("title", "Spreadsheet Data")
            document.add_heading(title, level=1)
            
            # Add table
            table = document.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = parameters.get("table_style", "Table Grid")
            
            # Add header row
            for col_idx, column in enumerate(df.columns):
                table.cell(0, col_idx).text = str(column)
            
            # Add data rows
            for row_idx, row in enumerate(df.itertuples(index=False)):
                for col_idx, value in enumerate(row):
                    table.cell(row_idx + 1, col_idx).text = str(value)
            
            # Save the document
            document.save(file_path)
        except ImportError:
            raise ConversionError(
                "python-docx is required for DOCX conversion. "
                "Install it with 'pip install python-docx'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to save DOCX file: {str(e)}")
    
    def _save_pdf(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a PDF file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        try:
            # First convert to HTML
            html_content = df.to_html(
                index=parameters.get("index", False),
                table_id=parameters.get("table_id", "data"),
                classes=parameters.get("classes", "dataframe"),
                escape=parameters.get("escape", True)
            )
            
            # Add CSS
            css_content = """
body { font-family: Arial, sans-serif; margin: 20px; }
table { border-collapse: collapse; width: 100%; margin-top: 20px; }
th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
th { background-color: #f2f2f2; }
.dataframe { margin-bottom: 20px; }
"""
            
            # Create full HTML document
            title = parameters.get("title", "Spreadsheet Data")
            html_doc = """<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>""" + title + """</title>
    <style>
""" + css_content + """
    </style>
</head>
<body>
    <h1>""" + title + """</h1>
""" + html_content + """
</body>
</html>"""
            
            # Convert HTML to PDF
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_html:
                temp_html_path = temp_html.name
                temp_html.write(html_doc.encode('utf-8'))
            
            try:
                # Try to use WeasyPrint
                import weasyprint
                weasyprint.HTML(filename=temp_html_path).write_pdf(
                    file_path,
                    stylesheets=[],
                    presentational_hints=True
                )
            except ImportError:
                try:
                    # Try to use pdfkit with wkhtmltopdf
                    import pdfkit
                    options = {
                        'page-size': parameters.get("page_size", "A4"),
                        'orientation': parameters.get("orientation", "portrait"),
                        'margin-top': f"{parameters.get('margin', 1.0)}in",
                        'margin-right': f"{parameters.get('margin', 1.0)}in",
                        'margin-bottom': f"{parameters.get('margin', 1.0)}in",
                        'margin-left': f"{parameters.get('margin', 1.0)}in",
                    }
                    pdfkit.from_file(temp_html_path, file_path, options=options)
                except ImportError:
                    # Try to use LibreOffice headless conversion
                    import subprocess
                    result = subprocess.run(
                        [
                            "libreoffice", "--headless", "--convert-to", "pdf",
                            "--outdir", str(file_path.parent),
                            temp_html_path
                        ],
                        capture_output=True,
                        text=True,
                        check=True,
                    )
                    # Move the output file to the desired location if needed
                    output_path = Path(temp_html_path).with_suffix('.pdf')
                    if output_path.name != file_path.name:
                        import shutil
                        shutil.move(output_path, file_path)
            finally:
                # Cleanup temporary file
                import os
                if os.path.exists(temp_html_path):
                    os.unlink(temp_html_path)
        except Exception as e:
            raise ConversionError(f"Failed to save PDF file: {str(e)}")
    
    def _save_txt(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a plain text file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        encoding = parameters.get("encoding", "utf-8")
        delimiter = parameters.get("delimiter", "\t")
        
        try:
            with open(file_path, "w", encoding=encoding) as f:
                # Write title if provided
                title = parameters.get("title")
                if title:
                    f.write(f"{title}\n\n")
                
                # Write headers
                headers = df.columns
                f.write(delimiter.join(str(h) for h in headers) + "\n")
                
                # Write separator line
                f.write(delimiter.join(["-" * len(str(h)) for h in headers]) + "\n")
                
                # Write data rows
                for _, row in df.iterrows():
                    f.write(delimiter.join(str(v) for v in row) + "\n")
        except Exception as e:
            raise ConversionError(f"Failed to save TXT file: {str(e)}")
    
    def _save_md(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a Markdown file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        encoding = parameters.get("encoding", "utf-8")
        
        try:
            with open(file_path, "w", encoding=encoding) as f:
                # Write title if provided
                title = parameters.get("title", "Spreadsheet Data")
                f.write(f"# {title}\n\n")
                
                # Write headers
                headers = df.columns
                f.write("| " + " | ".join(str(h) for h in headers) + " |\n")
                
                # Write separator
                f.write("| " + " | ".join(["---"] * len(headers)) + " |\n")
                
                # Write data rows
                for _, row in df.iterrows():
                    f.write("| " + " | ".join(str(v) for v in row) + " |\n")
        except Exception as e:
            raise ConversionError(f"Failed to save Markdown file: {str(e)}")
    
    def _save_yaml(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a YAML file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        try:
            import yaml
            import datetime
            
            # Convert DataFrame to list of dictionaries
            data = df.to_dict(orient='records')
            
            # Add metadata if needed
            if parameters.get("include_metadata", False):
                yaml_data = {
                    "metadata": {
                        "title": parameters.get("title", "Spreadsheet Data"),
                        "created": str(datetime.datetime.now()),
                        "columns": list(df.columns),
                        "rows": len(df)
                    },
                    "data": data
                }
            else:
                yaml_data = data
            
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                yaml.dump(
                    yaml_data, 
                    f, 
                    default_flow_style=parameters.get("default_flow_style", False),
                    sort_keys=parameters.get("sort_keys", False),
                    allow_unicode=True
                )
        except ImportError:
            raise ConversionError(
                "PyYAML is required for YAML conversion. "
                "Install it with 'pip install pyyaml'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to save YAML file: {str(e)}")
    
    def _save_ini(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to an INI file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        try:
            import configparser
            import datetime
            
            config = configparser.ConfigParser()
            
            # Create sections and properties
            for idx, row in df.iterrows():
                section_name = parameters.get("section_prefix", "item") + str(idx)
                config[section_name] = {}
                
                for col in df.columns:
                    # Convert all values to strings for INI file
                    config[section_name][str(col)] = str(row[col])
            
            # Add metadata section if requested
            if parameters.get("include_metadata", True):
                config["metadata"] = {
                    "title": parameters.get("title", "Spreadsheet Data"),
                    "created": str(datetime.datetime.now()),
                    "columns": str(list(df.columns)),
                    "rows": str(len(df))
                }
            
            # Write to file
            with open(file_path, 'w', encoding='utf-8') as f:
                config.write(f)
        except Exception as e:
            raise ConversionError(f"Failed to save INI file: {str(e)}")
    
    def _save_toml(
        self, 
        df: "pd.DataFrame", 
        file_path: Path, 
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a TOML file.
        
        Args:
            df: Pandas DataFrame to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the file cannot be saved.
        """
        try:
            import datetime
            
            # Try different TOML libraries
            try:
                import tomli_w
                has_tomli_w = True
            except ImportError:
                try:
                    import toml
                    has_tomli_w = False
                except ImportError:
                    raise ConversionError(
                        "toml or tomli_w is required for TOML conversion. "
                        "Install it with 'pip install toml' or 'pip install tomli_w'."
                    )
            
            # Convert DataFrame to list of dictionaries
            items = df.to_dict(orient='records')
            
            # Create TOML data structure
            toml_data = {
                "items": items
            }
            
            # Add metadata if requested
            if parameters.get("include_metadata", True):
                toml_data["metadata"] = {
                    "title": parameters.get("title", "Spreadsheet Data"),
                    "created": str(datetime.datetime.now()),
                    "columns": list(df.columns),
                    "rows": len(df)
                }
            
            # Write to file
            if has_tomli_w:
                # Using tomli_w (Python 3.11+)
                with open(file_path, "wb") as f:
                    tomli_w.dump(toml_data, f)
            else:
                # Using toml
                with open(file_path, "w", encoding="utf-8") as f:
                    toml.dump(toml_data, f)
        except Exception as e:
            raise ConversionError(f"Failed to save TOML file: {str(e)}")

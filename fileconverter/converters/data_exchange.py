"""
Data exchange format converters for FileConverter.

This module provides converters for data exchange formats, including:
- JSON (.json)
- XML (.xml)
- YAML (.yaml, .yml)
- INI (.ini)
- TOML (.toml)
- CSV (.csv)
- TSV (.tsv)
"""

import configparser
import csv
import io
import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from fileconverter.core.registry import BaseConverter
from fileconverter.utils.error_handling import ConversionError
from fileconverter.utils.file_utils import get_file_extension, guess_encoding
from fileconverter.utils.logging_utils import get_logger
from fileconverter.utils.validation import validate_file_path

logger = get_logger(__name__)

# Define supported formats
SUPPORTED_FORMATS = ["json", "xml", "yaml", "yml", "ini", "toml", "csv", "tsv",
                   "docx", "pdf", "txt", "html", "htm", "md", "xlsx", "xls"]


class DataExchangeConverter(BaseConverter):
    """Converter for data exchange formats."""
    
    @classmethod
    def get_input_formats(cls) -> List[str]:
        """Get the list of input formats supported by this converter."""
        return ["json", "xml", "yaml", "ini", "toml", "csv", "tsv",
                "docx", "pdf", "txt", "html", "htm", "md", "xlsx"]
    
    @classmethod
    def get_output_formats(cls) -> List[str]:
        """Get the list of output formats supported by this converter."""
        return ["json", "xml", "yaml", "ini", "toml", "csv", "tsv",
                "docx", "pdf", "txt", "html", "md", "xlsx"]
    
    @classmethod
    def get_format_extensions(cls, format_name: str) -> List[str]:
        """Get the list of file extensions for a specific format."""
        format_map = {
            "json": ["json"],
            "xml": ["xml"],
            "yaml": ["yaml", "yml"],
            "ini": ["ini", "conf", "cfg"],
            "toml": ["toml"],
            "csv": ["csv"],
            "tsv": ["tsv"],
            "docx": ["docx"],
            "pdf": ["pdf"],
            "txt": ["txt"],
            "html": ["html", "htm"],
            "md": ["md", "markdown"],
            "xlsx": ["xlsx"],
            "xls": ["xls"],
        }
        return format_map.get(format_name.lower(), [])
    
    def convert(
        self, 
        input_path: Union[str, Path], 
        output_path: Union[str, Path],
        temp_dir: Union[str, Path],
        parameters: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Convert a data file from one format to another.
        
        Args:
            input_path: Path to the input file.
            output_path: Path where the output file will be saved.
            temp_dir: Directory for temporary files.
            parameters: Conversion parameters.
        
        Returns:
            Dictionary with information about the conversion.
        
        Raises:
            ConversionError: If the conversion fails.
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        temp_dir = Path(temp_dir)
        
        # Validate paths
        validate_file_path(input_path, must_exist=True)
        
        # Get formats from file extensions
        input_ext = get_file_extension(input_path).lower()
        output_ext = get_file_extension(output_path).lower()
        
        # Normalize formats
        input_format = self._normalize_format(input_ext)
        output_format = self._normalize_format(output_ext)
        
        if not input_format:
            raise ConversionError(f"Unsupported input format: {input_ext}")
        
        if not output_format:
            raise ConversionError(f"Unsupported output format: {output_ext}")
        
        # Load data from input file
        try:
            data = self._load_data(input_path, input_format, parameters)
        except Exception as e:
            raise ConversionError(f"Failed to load {input_format} data: {str(e)}")
        
        # Save data to output file
        try:
            self._save_data(data, output_path, output_format, parameters)
        except Exception as e:
            raise ConversionError(f"Failed to save {output_format} data: {str(e)}")
            
        # Handle conversions that require special processing
        if (input_format in ["json", "xml", "yaml", "ini", "toml", "csv", "tsv"] and
            output_format in ["docx", "pdf", "html", "md"]):
            try:
                # Convert to document format
                self._convert_to_document_format(data, output_path, output_format, parameters)
            except Exception as e:
                raise ConversionError(f"Failed to convert to {output_format}: {str(e)}")
        
        return {
            "input_format": input_format,
            "output_format": output_format,
            "input_path": str(input_path),
            "output_path": str(output_path),
        }
    
    def get_parameters(self) -> Dict[str, Dict[str, Any]]:
        """Get the parameters supported by this converter."""
        return {
            "json": {
                "indent": {
                    "type": "integer",
                    "description": "Number of spaces for indentation",
                    "default": 2,
                    "min": 0,
                    "max": 8,
                },
                "sort_keys": {
                    "type": "boolean",
                    "description": "Whether to sort keys alphabetically",
                    "default": False,
                },
                "ensure_ascii": {
                    "type": "boolean",
                    "description": "Escape non-ASCII characters",
                    "default": False,
                },
            },
            "xml": {
                "pretty": {
                    "type": "boolean",
                    "description": "Whether to pretty-print the XML",
                    "default": True,
                },
                "root_element": {
                    "type": "string",
                    "description": "Name of the root element",
                    "default": "root",
                },
                "item_element": {
                    "type": "string",
                    "description": "Name of the item element for arrays",
                    "default": "item",
                },
            },
            "yaml": {
                "explicit_start": {
                    "type": "boolean",
                    "description": "Whether to include document start marker",
                    "default": True,
                },
                "default_flow_style": {
                    "type": "boolean",
                    "description": "Use flow style for collections",
                    "default": False,
                },
            },
            "csv": {
                "delimiter": {
                    "type": "string",
                    "description": "Field delimiter",
                    "default": ",",
                },
                "quotechar": {
                    "type": "string",
                    "description": "Quote character",
                    "default": '"',
                },
                "encoding": {
                    "type": "string",
                    "description": "File encoding",
                    "default": "utf-8",
                },
                "header": {
                    "type": "boolean",
                    "description": "Whether to write a header row",
                    "default": True,
                },
                "flatten": {
                    "type": "boolean",
                    "description": "Whether to flatten nested data",
                    "default": True,
                },
            },
        }
    
    def _normalize_format(self, extension: str) -> Optional[str]:
        """Normalize format name from file extension.
        
        Args:
            extension: File extension.
        
        Returns:
            Normalized format name or None if not supported.
        """
        if extension in ["yml", "yaml"]:
            return "yaml"
        elif extension in ["csv"]:
            return "csv"
        elif extension in ["tsv"]:
            return "tsv"
        elif extension in ["json", "xml", "ini", "toml"]:
            return extension
        
        return None
    
    def _load_data(
        self, 
        file_path: Path, 
        format_name: str,
        parameters: Dict[str, Any]
    ) -> Any:
        """Load data from a file.
        
        Args:
            file_path: Path to the file.
            format_name: Format of the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        if format_name == "json":
            return self._load_json(file_path, parameters)
        elif format_name == "xml":
            return self._load_xml(file_path, parameters)
        elif format_name == "yaml":
            return self._load_yaml(file_path, parameters)
        elif format_name == "ini":
            return self._load_ini(file_path, parameters)
        elif format_name == "toml":
            return self._load_toml(file_path, parameters)
        elif format_name == "csv":
            return self._load_csv(file_path, parameters)
        elif format_name == "tsv":
            return self._load_tsv(file_path, parameters)
        else:
            raise ConversionError(f"Unsupported format: {format_name}")
    
    def _save_data(
        self, 
        data: Any, 
        file_path: Path, 
        format_name: str,
        parameters: Dict[str, Any]
    ) -> None:
        """Save data to a file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            format_name: Format of the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        if format_name == "json":
            self._save_json(data, file_path, parameters)
        elif format_name == "xml":
            self._save_xml(data, file_path, parameters)
        elif format_name == "yaml":
            self._save_yaml(data, file_path, parameters)
        elif format_name == "ini":
            self._save_ini(data, file_path, parameters)
        elif format_name == "toml":
            self._save_toml(data, file_path, parameters)
        elif format_name == "csv":
            self._save_csv(data, file_path, parameters)
        elif format_name == "tsv":
            self._save_tsv(data, file_path, parameters)
        else:
            raise ConversionError(f"Unsupported format: {format_name}")
    
    def _load_json(self, file_path: Path, parameters: Dict[str, Any]) -> Any:
        """Load data from a JSON file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        encoding = parameters.get("encoding", "utf-8")
        
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return json.load(f)
        except UnicodeDecodeError:
            # Try to detect encoding
            encoding = guess_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                return json.load(f)
        except json.JSONDecodeError as e:
            raise ConversionError(f"Invalid JSON: {str(e)}")
        except Exception as e:
            raise ConversionError(f"Failed to load JSON: {str(e)}")
    
    def _load_xml(self, file_path: Path, parameters: Dict[str, Any]) -> Any:
        """Load data from an XML file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        try:
            import xmltodict
            
            with open(file_path, "rb") as f:
                return xmltodict.parse(f.read())
        except ImportError:
            raise ConversionError(
                "xmltodict is required for XML conversion. "
                "Install it with 'pip install xmltodict'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to load XML: {str(e)}")
    
    def _load_yaml(self, file_path: Path, parameters: Dict[str, Any]) -> Any:
        """Load data from a YAML file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        try:
            import yaml
            
            with open(file_path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f)
        except ImportError:
            raise ConversionError(
                "PyYAML is required for YAML conversion. "
                "Install it with 'pip install pyyaml'."
            )
        except yaml.YAMLError as e:
            raise ConversionError(f"Invalid YAML: {str(e)}")
        except Exception as e:
            raise ConversionError(f"Failed to load YAML: {str(e)}")
    
    def _load_ini(self, file_path: Path, parameters: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
        """Load data from an INI file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        try:
            config = configparser.ConfigParser()
            config.read(file_path)
            
            # Convert to dictionary
            result = {}
            for section in config.sections():
                result[section] = {}
                for key, value in config[section].items():
                    result[section][key] = value
            
            return result
        except Exception as e:
            raise ConversionError(f"Failed to load INI: {str(e)}")
    
    def _load_toml(self, file_path: Path, parameters: Dict[str, Any]) -> Any:
        """Load data from a TOML file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        try:
            import tomli  # Python 3.11+ built-in
            
            with open(file_path, "rb") as f:
                return tomli.load(f)
        except ImportError:
            try:
                import toml
                
                with open(file_path, "r", encoding="utf-8") as f:
                    return toml.load(f)
            except ImportError:
                raise ConversionError(
                    "toml or tomli is required for TOML conversion. "
                    "Install it with 'pip install toml'."
                )
            except Exception as e:
                raise ConversionError(f"Failed to load TOML: {str(e)}")
        except Exception as e:
            raise ConversionError(f"Failed to load TOML: {str(e)}")
    
    def _load_csv(self, file_path: Path, parameters: Dict[str, Any]) -> List[Dict[str, str]]:
        """Load data from a CSV file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data as a list of dictionaries.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        delimiter = parameters.get("delimiter", ",")
        quotechar = parameters.get("quotechar", '"')
        encoding = parameters.get("encoding")
        
        if not encoding:
            encoding = guess_encoding(file_path)
        
        try:
            result = []
            with open(file_path, "r", encoding=encoding, newline="") as f:
                reader = csv.DictReader(f, delimiter=delimiter, quotechar=quotechar)
                for row in reader:
                    result.append(dict(row))
            
            return result
        except Exception as e:
            raise ConversionError(f"Failed to load CSV: {str(e)}")
    
    def _load_tsv(self, file_path: Path, parameters: Dict[str, Any]) -> List[Dict[str, str]]:
        """Load data from a TSV file.
        
        Args:
            file_path: Path to the file.
            parameters: Conversion parameters.
        
        Returns:
            Loaded data as a list of dictionaries.
        
        Raises:
            ConversionError: If the data cannot be loaded.
        """
        # Use CSV loader with tab delimiter
        parameters["delimiter"] = "\t"
        return self._load_csv(file_path, parameters)
    
    def _save_json(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to a JSON file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        indent = parameters.get("indent", 2)
        sort_keys = parameters.get("sort_keys", False)
        ensure_ascii = parameters.get("ensure_ascii", False)
        encoding = parameters.get("encoding", "utf-8")
        
        try:
            with open(file_path, "w", encoding=encoding) as f:
                json.dump(
                    data,
                    f,
                    indent=indent,
                    sort_keys=sort_keys,
                    ensure_ascii=ensure_ascii
                )
        except Exception as e:
            raise ConversionError(f"Failed to save JSON: {str(e)}")
    
    def _save_xml(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to an XML file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        try:
            import xmltodict
            from dicttoxml import dicttoxml
            
            pretty = parameters.get("pretty", True)
            root_element = parameters.get("root_element", "root")
            item_element = parameters.get("item_element", "item")
            encoding = parameters.get("encoding", "utf-8")
            
            # If data is already in XML format (from xmltodict.parse())
            if isinstance(data, dict) and len(data) == 1:
                # Try to use existing root element
                xml_data = xmltodict.unparse(
                    data,
                    pretty=pretty,
                    encoding=encoding
                )
            else:
                # Convert to XML with dicttoxml
                xml_data = dicttoxml(
                    data,
                    custom_root=root_element,
                    item_func=lambda x: item_element,
                    attr_type=False
                )
                
                # Pretty print if requested
                if pretty:
                    try:
                        from lxml import etree
                        root = etree.fromstring(xml_data)
                        xml_data = etree.tostring(
                            root,
                            pretty_print=True,
                            encoding=encoding
                        )
                    except ImportError:
                        logger.warning("lxml not available, XML will not be pretty-printed")
            
            # Write to file
            with open(file_path, "wb") as f:
                f.write(xml_data)
        
        except ImportError:
            raise ConversionError(
                "xmltodict and dicttoxml are required for XML conversion. "
                "Install them with 'pip install xmltodict dicttoxml'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to save XML: {str(e)}")
    
    def _save_yaml(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to a YAML file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        try:
            import yaml
            
            explicit_start = parameters.get("explicit_start", True)
            default_flow_style = parameters.get("default_flow_style", False)
            encoding = parameters.get("encoding", "utf-8")
            
            with open(file_path, "w", encoding=encoding) as f:
                yaml.dump(
                    data,
                    f,
                    explicit_start=explicit_start,
                    default_flow_style=default_flow_style
                )
        except ImportError:
            raise ConversionError(
                "PyYAML is required for YAML conversion. "
                "Install it with 'pip install pyyaml'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to save YAML: {str(e)}")
    
    def _save_ini(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to an INI file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        try:
            # INI files require a specific structure
            if not isinstance(data, dict):
                raise ConversionError(
                    "Invalid data for INI format. Expected dictionary with sections."
                )
            
            config = configparser.ConfigParser()
            
            # Add sections and values
            for section, section_data in data.items():
                if not isinstance(section_data, dict):
                    raise ConversionError(
                        f"Invalid data for INI section '{section}'. Expected dictionary."
                    )
                
                config[section] = {}
                for key, value in section_data.items():
                    # Convert values to strings
                    if value is not None:
                        config[section][key] = str(value)
            
            # Write to file
            with open(file_path, "w", encoding="utf-8") as f:
                config.write(f)
        
        except Exception as e:
            raise ConversionError(f"Failed to save INI: {str(e)}")
    
    def _save_toml(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to a TOML file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        try:
            try:
                import tomli_w  # Python 3.11+ built-in
                
                with open(file_path, "wb") as f:
                    tomli_w.dump(data, f)
            except ImportError:
                import toml
                
                with open(file_path, "w", encoding="utf-8") as f:
                    toml.dump(data, f)
        except ImportError:
            raise ConversionError(
                "toml/tomli_w is required for TOML conversion. "
                "Install it with 'pip install toml'."
            )
        except Exception as e:
            raise ConversionError(f"Failed to save TOML: {str(e)}")
    
    def _save_csv(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to a CSV file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        delimiter = parameters.get("delimiter", ",")
        quotechar = parameters.get("quotechar", '"')
        encoding = parameters.get("encoding", "utf-8")
        header = parameters.get("header", True)
        flatten = parameters.get("flatten", True)
        
        try:
            # Ensure data is in a list of dictionaries format
            rows = self._normalize_data_for_csv(data, flatten)
            
            if not rows:
                raise ConversionError("No data to write to CSV")
            
            # Get fieldnames from the first row
            fieldnames = list(rows[0].keys())
            
            with open(file_path, "w", encoding=encoding, newline="") as f:
                writer = csv.DictWriter(
                    f,
                    fieldnames=fieldnames,
                    delimiter=delimiter,
                    quotechar=quotechar,
                    quoting=csv.QUOTE_MINIMAL
                )
                
                if header:
                    writer.writeheader()
                
                for row in rows:
                    writer.writerow(row)
        
        except Exception as e:
            raise ConversionError(f"Failed to save CSV: {str(e)}")
    
    def _save_tsv(self, data: Any, file_path: Path, parameters: Dict[str, Any]) -> None:
        """Save data to a TSV file.
        
        Args:
            data: Data to save.
            file_path: Path where to save the file.
            parameters: Conversion parameters.
        
        Raises:
            ConversionError: If the data cannot be saved.
        """
        # Use CSV saver with tab delimiter
        parameters["delimiter"] = "\t"
        return self._save_csv(data, file_path, parameters)
    
    def _convert_to_document_format(
        self,
        data: Any,
        output_path: Path,
        output_format: str,
        parameters: Dict[str, Any]
    ) -> None:
        """Convert structured data to a document format.
        
        Args:
            data: Data to convert.
            output_path: Path where to save the file.
            output_format: Output format.
            parameters: Conversion parameters.
            
        Raises:
            ConversionError: If the conversion fails.
        """
        try:
            # Convert data to a pandas DataFrame for easier handling
            import pandas as pd
            
            # Handle different data types
            if isinstance(data, list):
                df = pd.DataFrame(data)
            elif isinstance(data, dict):
                if any(isinstance(data[k], (dict, list)) for k in data):
                    # Complex nested structure
                    df = pd.json_normalize(data)
                else:
                    # Simple flat dictionary
                    df = pd.DataFrame([data])
            else:
                # Try to convert to DataFrame
                df = pd.DataFrame(data)
            
            # Now output to the target document format
            if output_format == "html":
                # Generate HTML table
                html_table = df.to_html(
                    index=parameters.get("index", False),
                    table_id=parameters.get("table_id", "data"),
                    classes=parameters.get("classes", "dataframe"),
                    escape=parameters.get("escape", True)
                )
                
                # Create full HTML document
                title = parameters.get("title", "Data Exchange")
                css_content = """
                body { font-family: Arial, sans-serif; margin: 20px; }
                table { border-collapse: collapse; width: 100%; margin-top: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .dataframe { margin-bottom: 20px; }
                """
                
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <title>{title}</title>
                    <style>
                        {css_content}
                    </style>
                </head>
                <body>
                    <h1>{title}</h1>
                    {html_table}
                </body>
                </html>
                """
                
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                    
            elif output_format == "pdf":
                # Convert to PDF via HTML
                import tempfile
                
                # First create HTML
                html_path = tempfile.NamedTemporaryFile(suffix=".html", delete=False).name
                
                # Generate HTML content
                html_table = df.to_html(
                    index=parameters.get("index", False),
                    table_id=parameters.get("table_id", "data"),
                    classes=parameters.get("classes", "dataframe"),
                    escape=parameters.get("escape", True)
                )
                
                title = parameters.get("title", "Data Exchange")
                css_content = """
                body { font-family: Arial, sans-serif; margin: 20px; }
                table { border-collapse: collapse; width: 100%; margin-top: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .dataframe { margin-bottom: 20px; }
                """
                
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1">
                    <title>{title}</title>
                    <style>
                        {css_content}
                    </style>
                </head>
                <body>
                    <h1>{title}</h1>
                    {html_table}
                </body>
                </html>
                """
                
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                
                try:
                    # Try WeasyPrint
                    import weasyprint
                    weasyprint.HTML(filename=html_path).write_pdf(output_path)
                except ImportError:
                    try:
                        # Try pdfkit
                        import pdfkit
                        options = {
                            'page-size': parameters.get("page_size", "A4"),
                            'orientation': parameters.get("orientation", "portrait"),
                            'margin-top': f"{parameters.get('margin', 1.0)}in",
                            'margin-right': f"{parameters.get('margin', 1.0)}in",
                            'margin-bottom': f"{parameters.get('margin', 1.0)}in",
                            'margin-left': f"{parameters.get('margin', 1.0)}in",
                        }
                        pdfkit.from_file(html_path, str(output_path), options=options)
                    except ImportError:
                        raise ConversionError(
                            "Either WeasyPrint or pdfkit is required for PDF conversion. "
                            "Install with 'pip install weasyprint' or 'pip install pdfkit'."
                        )
                finally:
                    # Cleanup temporary file
                    import os
                    if os.path.exists(html_path):
                        os.unlink(html_path)
                        
            elif output_format == "docx":
                # Convert to DOCX
                try:
                    import docx
                    from docx import Document
                    
                    doc = Document()
                    
                    # Add title
                    title = parameters.get("title", "Data Exchange")
                    doc.add_heading(title, level=1)
                    
                    # Add table
                    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
                    table.style = parameters.get("table_style", "Table Grid")
                    
                    # Add header row
                    for col_idx, column in enumerate(df.columns):
                        table.cell(0, col_idx).text = str(column)
                    
                    # Add data rows
                    for row_idx, row in enumerate(df.iterrows()):
                        for col_idx, value in enumerate(row[1]):
                            table.cell(row_idx + 1, col_idx).text = str(value)
                    
                    # Save document
                    doc.save(output_path)
                except ImportError:
                    raise ConversionError(
                        "python-docx is required for DOCX conversion. "
                        "Install with 'pip install python-docx'."
                    )
                    
            elif output_format == "md":
                # Convert to Markdown
                with open(output_path, "w", encoding="utf-8") as f:
                    # Write title
                    title = parameters.get("title", "Data Exchange")
                    f.write(f"# {title}\n\n")
                    
                    # Write table header
                    f.write("| " + " | ".join(str(col) for col in df.columns) + " |\n")
                    f.write("| " + " | ".join(["---"] * len(df.columns)) + " |\n")
                    
                    # Write rows
                    for _, row in df.iterrows():
                        f.write("| " + " | ".join(str(val) for val in row) + " |\n")
        
        except Exception as e:
            raise ConversionError(f"Failed to convert to {output_format}: {str(e)}")
            
    def _normalize_data_for_csv(
        self, 
        data: Any, 
        flatten: bool = True
    ) -> List[Dict[str, str]]:
        """Normalize data for CSV format.
        
        Args:
            data: Data to normalize.
            flatten: Whether to flatten nested data.
        
        Returns:
            Normalized data as a list of dictionaries.
        
        Raises:
            ConversionError: If the data cannot be normalized.
        """
        # If data is a dictionary, wrap it in a list
        if isinstance(data, dict):
            if all(isinstance(v, dict) for v in data.values()):
                # Dictionary of dictionaries, convert to list of dictionaries
                rows = [{"name": k, **v} for k, v in data.items()]
            else:
                # Single dictionary, make it a row
                rows = [data]
        elif isinstance(data, list):
            rows = data
        else:
            raise ConversionError(
                f"Invalid data type for CSV: {type(data)}. "
                "Expected list or dictionary."
            )
        
        # Flatten nested data if requested
        if flatten:
            flattened_rows = []
            for row in rows:
                if not isinstance(row, dict):
                    # Skip non-dictionary rows
                    continue
                
                flat_row = {}
                for key, value in row.items():
                    if isinstance(value, dict):
                        # Flatten nested dictionaries
                        for subkey, subvalue in value.items():
                            flat_key = f"{key}.{subkey}"
                            flat_row[flat_key] = subvalue
                    elif isinstance(value, list):
                        # Convert lists to comma-separated strings
                        flat_row[key] = ", ".join(str(item) for item in value)
                    else:
                        # Keep simple values as is
                        flat_row[key] = value
                
                flattened_rows.append(flat_row)
            
            rows = flattened_rows
        
        # Convert all values to strings
        for row in rows:
            for key, value in row.items():
                if value is None:
                    row[key] = ""
                elif not isinstance(value, str):
                    row[key] = str(value)
        
        return rows
